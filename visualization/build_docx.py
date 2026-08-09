# -*- coding: utf-8 -*-
"""Build master_thesis/магистерски_труд.docx from the markdown parts,
following the official FEIT урнек:

- A4, urnek margins, footer page numbers (front matter roman, body arabic from 1)
- real Heading 1-3 styles (Arial per урнек) with automatic multilevel numbering
  (literal "7.3."-style numbers are stripped from the markdown at build time)
- Содржина / Листа на слики / Листа на табели as live Word fields
  (<!--toc-->, <!--lof-->, <!--lot--> markers in part0.md)
- figure captions below/centered, table captions above/left, both in the
  Caption style with live SEQ numbering
- citation footnotes (^[...]) converted to numbered [n] references matched
  against the bibliography in part6.md; unmatched (explanatory) footnotes stay

Requires: pip install pypandoc-binary python-docx
Run the figure scripts first (see README.md in this folder).
Open the result in Word and press Ctrl+A, F9 to fill in all fields.
"""
import os
import re
import pypandoc

# --- CONFIG -----------------------------------------------------------
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(REPO, "master_thesis", "src")   # thesis text + figures
OUT = os.path.join(REPO, "master_thesis", "магистерски_труд.docx")
# ----------------------------------------------------------------------
PARTS = [os.path.join(SRC, "thesis", f"part{i}.md") for i in range(7)]
COMBINED = os.path.join(SRC, "combined.md")

PAGEBREAK = '\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n'


def field_block(instr, placeholder):
    return ('\n```{=openxml}\n'
            '<w:p><w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
            f'<w:r><w:instrText xml:space="preserve"> {instr} </w:instrText></w:r>'
            '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
            f'<w:r><w:t>{placeholder}</w:t></w:r>'
            '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>\n```\n')


text = []
for p in PARTS:
    with open(p, encoding="utf-8") as f:
        text.append(f.read())
doc = "\n\n".join(text)

# ---- 1. citation footnotes -> [n] references ------------------------------
bib = dict(re.findall(r'(?m)^\[(\d+)\]\s+(.+)$', doc))


def _norm(s):
    return re.sub(r'[^a-z0-9а-ш]+', ' ', s.lower())


def match_entry(note):
    arx = re.search(r'(\d{4}\.\d{4,5})', note)
    if arx:
        hits = [n for n, e in bib.items() if arx.group(1) in e]
        if len(hits) == 1:
            return hits[0]
    title = re.search(r'[""“"]([^""”"]{8,})[""”"]', note)
    if title:
        t = _norm(title.group(1))
        hits = [n for n, e in bib.items() if t in _norm(e)]
        if len(hits) == 1:
            return hits[0]
    surname = re.match(r'^(?:[A-Z]\.\s*)*([A-Z][a-zА-Я][\w-]+)', note.strip())
    if surname:
        s = surname.group(1)
        hits = [n for n, e in bib.items() if s in e.split('*')[0] or s in e[:120]]
        if len(hits) == 1:
            return hits[0]
    # capitalized multi-word phrases (venue or partial title), longest first
    phrases = sorted(re.findall(r'[A-Z][a-z]+(?: [A-Z][a-z]+)+', note),
                     key=len, reverse=True)
    for ph in phrases:
        hits = [n for n, e in bib.items() if ph in e]
        if len(hits) == 1:
            return hits[0]
    words = [w for w in re.findall(r'[A-Z][a-z]{3,}', note)][:4]
    for w in words:
        hits = [n for n, e in bib.items() if w in e[:130]]
        if len(hits) == 1:
            return hits[0]
    return None


converted, kept = 0, []
def replace_note(m):
    global converted
    n = match_entry(m.group(1))
    if n is None:
        kept.append(m.group(1)[:60])
        return m.group(0)
    converted += 1
    return f'[{n}]'

doc = re.sub(r'\^\[([^\]]+)\]', replace_note, doc)
print(f'footnotes -> [n]: {converted} converted, {len(kept)} kept as footnotes')
for k in kept:
    print('  kept:', k)

# ---- 2. strip literal heading numbers (auto-numbering takes over) ---------
doc = re.sub(r'(?m)^(#{1,4}) (\d+(?:\.\d+)*)\.\s+', r'\1 ', doc)

# ---- 3. markers -----------------------------------------------------------
doc = doc.replace("<!--pagebreak-->", PAGEBREAK)
doc = doc.replace("<!--toc-->", field_block(r'TOC \o "1-3" \h \z \u',
                                            'Селектирај и притисни F9 за содржина.'))
doc = doc.replace("<!--lof-->", field_block(r'TOC \h \z \c "Слика"',
                                            'Селектирај и притисни F9 за листа на слики.'))
doc = doc.replace("<!--lot-->", field_block(r'TOC \h \z \c "Табела"',
                                            'Селектирај и притисни F9 за листа на табели.'))

missing = [line.split("](")[1].split(")")[0]
           for line in doc.splitlines()
           if line.startswith("![") and "](" in line
           and not os.path.exists(os.path.join(SRC, line.split("](")[1].split(")")[0]))]
assert not missing, f"missing figures: {missing} — run the make_*.py scripts first"

with open(COMBINED, "w", encoding="utf-8") as f:
    f.write(doc)

pypandoc.convert_file(
    COMBINED, "docx", outputfile=OUT,
    format="markdown+footnotes+pipe_tables+raw_attribute+tex_math_dollars",
    extra_args=["--metadata", "lang=mk", f"--resource-path={SRC}"],
)
print("pandoc done ->", OUT)

# ---- restyle with python-docx (урнек styles + numbering + sections) -------
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

d = docx.Document(OUT)


def set_font(style, name=None, size=None, color=None, bold=None, italic=None):
    f = style.font
    if name:
        f.name = name
        rpr = style.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rf)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(attr), name)
    if size: f.size = Pt(size)
    if color is not None: f.color.rgb = RGBColor(*color)
    if bold is not None: f.bold = bold
    if italic is not None: f.italic = italic


BODY = "Times New Roman"
HEAD = "Arial"          # урнек: naslovite se Arial
styles = d.styles
wanted = {
    "Normal":          dict(name=BODY, size=12),
    "Body Text":       dict(name=BODY, size=12),
    "First Paragraph": dict(name=BODY, size=12),
    "Compact":         dict(name=BODY, size=11),
    "Title":           dict(name=BODY, size=24, color=(0, 0, 0), bold=True),
    "Heading 1":       dict(name=HEAD, size=14, color=(0, 0, 0), bold=True),
    "Heading 2":       dict(name=HEAD, size=12, color=(0, 0, 0), bold=True),
    "Heading 3":       dict(name=HEAD, size=12, color=(0, 0, 0), bold=False),
    "Heading 4":       dict(name=BODY, size=12, color=(0, 0, 0), bold=True, italic=False),
    "Image Caption":   dict(name=BODY, size=10, color=(0, 0, 0), bold=True, italic=False),
    "Table Caption":   dict(name=BODY, size=10, color=(0, 0, 0), bold=True, italic=False),
    "Caption":         dict(name=BODY, size=10, color=(0, 0, 0), bold=True, italic=False),
    "Footnote Text":   dict(name=BODY, size=8),
    "Block Text":      dict(name=BODY, size=11),
}
for sname, kw in wanted.items():
    try:
        set_font(styles[sname], **kw)
    except KeyError:
        print("style not found:", sname)

# урнек: Normal 12 pt, edinicен prored, justified
for sname in ("Normal", "Body Text", "First Paragraph"):
    try:
        pf = styles[sname].paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(6)
    except KeyError:
        pass
try:
    pf = styles["Footnote Text"].paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
except KeyError:
    pass

# ---- automatic heading numbering ------------------------------------------
numbering = d.part.numbering_part.element
ABSTRACT_ID, NUM_ID = 90, 90
lvls = ''.join(
    f'<w:lvl w:ilvl="{i}"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
    f'<w:lvlText w:val="{lvltext}"/><w:suff w:val="space"/>'
    f'<w:lvlJc w:val="left"/><w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr></w:lvl>'
    for i, lvltext in enumerate(["%1.", "%1.%2.", "%1.%2.%3.", "%1.%2.%3.%4."]))
abstract = parse_xml(f'<w:abstractNum {W} w:abstractNumId="{ABSTRACT_ID}">'
                     f'<w:multiLevelType w:val="multilevel"/>{lvls}</w:abstractNum>')
first_num = numbering.find(qn('w:num'))
numbering.insert(list(numbering).index(first_num) if first_num is not None else 0, abstract)
numbering.append(parse_xml(
    f'<w:num {W} w:numId="{NUM_ID}"><w:abstractNumId w:val="{ABSTRACT_ID}"/></w:num>'))

for ilvl, sname in enumerate(["Heading 1", "Heading 2", "Heading 3", "Heading 4"]):
    st = styles[sname].element
    ppr = st.get_or_add_pPr()
    ppr.append(parse_xml(f'<w:numPr {W}><w:ilvl w:val="{ilvl}"/>'
                         f'<w:numId w:val="{NUM_ID}"/></w:numPr>'))

UNNUMBERED = ("Содржина", "Листа на слики", "Листа на табели", "Апстракт",
              "Abstract", "Листа на кратенки", "Користена литература",
              "Прилози", "Прилог ", "Примена на трансформер")
HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}


def suppress_number(par):
    ppr = par._p.get_or_add_pPr()
    ppr.append(parse_xml(f'<w:numPr {W}><w:ilvl w:val="0"/>'
                         f'<w:numId w:val="0"/></w:numPr>'))


for par in d.paragraphs:
    if par.style.name in HEADING_STYLES and \
            any(par.text.strip().startswith(u) for u in UNNUMBERED):
        suppress_number(par)

# ---- captions: Caption style + live SEQ numbers ---------------------------
def seq_field(label):
    return [parse_xml(f'<w:r {W}><w:fldChar w:fldCharType="begin"/></w:r>'),
            parse_xml(f'<w:r {W}><w:instrText xml:space="preserve">'
                      f' SEQ {label} \\* ARABIC </w:instrText></w:r>'),
            parse_xml(f'<w:r {W}><w:fldChar w:fldCharType="separate"/></w:r>'),
            parse_xml(f'<w:r {W}><w:t>0</w:t></w:r>'),
            parse_xml(f'<w:r {W}><w:fldChar w:fldCharType="end"/></w:r>')]


def rebuild_caption(par, label, tail, align):
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    try:
        par.style = styles["Caption"]
    except KeyError:
        pass
    par.alignment = align
    par.add_run(f"{label} ")
    for el in seq_field(label):
        par._p.append(el)
    par.add_run(f". {tail}" if tail else ".")


n_fig = n_tab = 0
for par in d.paragraphs:
    txt = par.text.strip()
    m = re.match(r'^Слика (\d+)\.\s*(.*)$', txt)
    if m and par.style.name in ("Image Caption", "Caption"):
        rebuild_caption(par, "Слика", m.group(2), WD_ALIGN_PARAGRAPH.CENTER)
        n_fig += 1
        continue
    m = re.match(r'^Табела (\d+)\.\s*(.*)$', txt)
    if m and par.style.name not in HEADING_STYLES:
        rebuild_caption(par, "Табела", m.group(2), WD_ALIGN_PARAGRAPH.LEFT)
        n_tab += 1
print(f"captions with SEQ fields: {n_fig} figures, {n_tab} tables")

# ---- sections: front matter (roman) | body (arabic from 1) ---------------
def page_setup(sectpr_xml_extra=""):
    return (f'<w:sectPr {W}><w:pgSz w:w="11906" w:h="16838"/>'
            '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"'
            ' w:header="709" w:footer="709" w:gutter="0"/>'
            f'{sectpr_xml_extra}</w:sectPr>')


body_start = None
for par in d.paragraphs:
    if par.style.name == "Heading 1" and par.text.strip() == "Вовед":
        body_start = par
        break
assert body_start is not None, "Вовед heading not found"

# a paragraph *before* Вовед carries the end-of-front-matter section break
brk = body_start.insert_paragraph_before("")
brk._p.get_or_add_pPr().append(
    parse_xml(page_setup('<w:pgNumType w:fmt="lowerRoman" w:start="1"/>')))

# document-level (final) section: A4 + arabic numbering from 1
final_sect = d.sections[-1]
final_sect.page_width, final_sect.page_height = Cm(21.0), Cm(29.7)
for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(final_sect, attr, Cm(2.54))
spr = final_sect._sectPr
pgnum = spr.find(qn('w:pgNumType'))
if pgnum is None:
    pgnum = parse_xml(f'<w:pgNumType {W}/>')
    spr.append(pgnum)
pgnum.set(qn('w:start'), '1')
pgnum.set(qn('w:fmt'), 'decimal')

d.save(OUT)

# reopen: now two sections exist; give them a shared centered PAGE footer
d = docx.Document(OUT)
first = d.sections[0]
first.page_width, first.page_height = Cm(21.0), Cm(29.7)
for attr in ("top_margin", "bottom_margin", "right_margin"):
    setattr(first, attr, Cm(2.54))
first.left_margin = Cm(2.54)

# насловната страница без број (урнек/правилник): different first page, празен footer
first.different_first_page_header_footer = True

footer = first.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for el in [parse_xml(f'<w:r {W}><w:fldChar w:fldCharType="begin"/></w:r>'),
           parse_xml(f'<w:r {W}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'),
           parse_xml(f'<w:r {W}><w:fldChar w:fldCharType="end"/></w:r>')]:
    fp._p.append(el)
for sect in d.sections[1:]:
    sect.footer.is_linked_to_previous = True

d.save(OUT)
print("styling + numbering + sections done")

# ---- quick verification ---------------------------------------------------
import zipfile
z = zipfile.ZipFile(OUT)
names = z.namelist()
media = [n for n in names if n.startswith("word/media/")]
has_footnotes = "word/footnotes.xml" in names
fx = z.read("word/footnotes.xml").decode("utf-8") if has_footnotes else ""
n_footnotes = fx.count("<w:footnote ") - 2 if has_footnotes else 0
docxml = z.read("word/document.xml").decode("utf-8")
numxml = z.read("word/numbering.xml").decode("utf-8")
print(f"images embedded: {len(media)}")
print(f"remaining true footnotes: {n_footnotes}")
print(f"page breaks: {docxml.count('w:br w:type=\"page\"')}")
print(f"math blocks: {docxml.count('<m:oMath')}")
print(f"TOC/list fields: {docxml.count('instrText')}, SEQ fields: {docxml.count('SEQ ')}")
print(f"heading num attached: {'abstractNumId w:val=\"90\"' in numxml}")
print(f"sections: {docxml.count('<w:sectPr')}")
print(f"size: {os.path.getsize(OUT)/1e6:.1f} MB")
print("NOTE: open in Word, Ctrl+A then F9 to populate TOC/lists/caption numbers.")
